# this file handles loading and processing PDFs into searchable chunks
import logging
import os, shutil
from langchain_community.document_loaders import PyPDFLoader
from app.rag.chunking import chunk_documents
from app.vectorstore.faiss_store import save_vectorstore
from fastapi import UploadFile

logger = logging.getLogger(__name__)
# folder where all uploaded PDFs are stored
UPLOAD_DIR = "app/data/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# takes a PDF file and processes it into searchable chunks stored in our database
def ingest_pdf(input_source):
    # figure out if we got an uploaded file or a file path
    if isinstance(input_source, UploadFile):
        # if its an upload, save it to our uploads folder first
        filename = input_source.filename
        persistent_path = os.path.join(UPLOAD_DIR, filename)
        with open(persistent_path, "wb") as f:
            shutil.copyfileobj(input_source.file, f)

    elif isinstance(input_source, str):
        # if its a file path, just use it directly
        persistent_path = input_source
        filename = os.path.basename(persistent_path)
        if not os.path.exists(persistent_path):
            raise FileNotFoundError(f"PDF not found: {persistent_path}")
    else:
        raise TypeError("ingest_pdf expects UploadFile or file path")

    try:
        logger.info(f"Loading PDF: {filename}")

        # load the PDF based on its file type
        if persistent_path.lower().endswith(".pdf"):
            # use PyPDF to extract text from each page
            loader = PyPDFLoader(persistent_path)
            documents = loader.load()
        elif persistent_path.lower().endswith(".docx"):
            # for DOCX files, read all paragraphs and combine them
            from docx import Document as DocxDocument
            from langchain_core.documents import Document
            doc = DocxDocument(persistent_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            documents = [Document(page_content="\n".join(full_text), metadata={"source": persistent_path})]
        else:
            raise ValueError("Unsupported file format")

        # make sure we actually got some text from the file
        if not documents:
            raise ValueError("PDF file is empty or unreadable")

        total_pages = len(documents)
        logger.info(f"Loaded {total_pages} pages from {filename}")

        # split the documents into smaller chunks for better search results
        chunks = chunk_documents(documents)
        logger.info(f"Created {len(chunks)} chunks")

        # save the chunks to our vector database so they can be searched later
        save_vectorstore(chunks)
        logger.info(f"Successfully ingested {filename}")

        # return info about what we processed
        return {
            "status": "success",
            "filename": filename,
            "pages": total_pages,
            "chunks": len(chunks),
        }

    except Exception as e:
        logger.error(f"Error processing PDF {filename}: {str(e)}")
        raise

# rebuilds the entire vector database from all remaining PDFs
# this is called after deleting a PDF to keep the database accurate
def rebuild_vectorstore_from_uploads():
    from app.vectorstore.faiss_store import replace_vectorstore
    from app.rag.chunking import chunk_documents
    from app.core.config import settings
    from langchain_community.document_loaders import PyPDFLoader
    import os
    import shutil

    uploads_dir = "app/data/uploads"

    # if no uploads folder exists, nothing to rebuild
    if not os.path.exists(uploads_dir):
        return

    # find all PDF and DOCX files in the uploads folder
    pdf_files = [
        f for f in os.listdir(uploads_dir)
        if f.lower().endswith(".pdf") or f.lower().endswith(".docx")
    ]

    # if all PDFs are deleted, clear the database and stop
    if not pdf_files:
        if os.path.exists(settings.VECTOR_DB_PATH):
            shutil.rmtree(settings.VECTOR_DB_PATH)
        logger.info("No PDFs remaining, vectorstore cleared")
        return

    documents = []

    # reload every remaining PDF
    for filename in pdf_files:
        try:
            if filename.lower().endswith(".pdf"):
                loader = PyPDFLoader(os.path.join(uploads_dir, filename))
                documents.extend(loader.load())
            elif filename.lower().endswith(".docx"):
                # handle DOCX files during rebuild too
                from docx import Document as DocxDocument
                from langchain_core.documents import Document
                path = os.path.join(uploads_dir, filename)
                doc = DocxDocument(path)
                text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                documents.append(Document(page_content=text, metadata={"source": path}))
            logger.info(f"Loaded {filename} for rebuild")
        except Exception as e:
            logger.warning(f"Failed to load {filename}: {e}")

    # if nothing could be loaded, clear the database
    if not documents:
        if os.path.exists(settings.VECTOR_DB_PATH):
            shutil.rmtree(settings.VECTOR_DB_PATH)
        return

    # split all documents into chunks and create a fresh database
    chunks = chunk_documents(documents)
    replace_vectorstore(chunks)
    logger.info(f"Rebuilt vectorstore with {len(chunks)} chunks from {len(pdf_files)} PDFs")
