# this service parses syllabus PDF/DOCX files and extracts structured data
# it finds the subject code, unit names, and topics from the document

import tempfile
import re
from typing import List, Dict, Optional, Tuple
from pypdf import PdfReader
from docx import Document


# reads a PDF file and pulls out the text and any table-like structures
def extract_text_from_pdf(path: str) -> Tuple[str, List[List[List[str]]]]:
    reader = PdfReader(path)
    text_parts = []
    tables = []
    
    # go through each page of the PDF
    for page in reader.pages:
        page_text = page.extract_text() or ""
        text_parts.append(page_text)
        
        # try to find tables by looking for pipe or tab separated content
        lines = page_text.split('\n')
        current_table = []
        
        for line in lines:
            # if line has pipes or tabs, it might be a table row
            if '|' in line or '\t' in line:
                cells = re.split(r'\s*\|\s*|\t+', line.strip())
                cells = [c.strip() for c in cells if c.strip()]
                if len(cells) >= 2:
                    current_table.append(cells)
            elif current_table:
                # end of table block, save if it has enough rows
                if len(current_table) >= 2:
                    tables.append(current_table)
                current_table = []
        
        # save any remaining table at end of page
        if current_table and len(current_table) >= 2:
            tables.append(current_table)
    
    return '\n'.join(text_parts), tables


# reads a DOCX file and extracts text paragraphs and actual tables
def extract_tables_from_docx(path: str) -> Tuple[str, List[List[List[str]]]]:
    doc = Document(path)
    text_parts = []
    tables = []
    
    # grab all the regular paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    
    # extract actual tables from the DOCX (much more reliable than PDF)
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_cells.append(cell_text)
            # skip completely empty rows
            if any(c for c in row_cells):
                table_data.append(row_cells)
        
        # need at least a header row and one data row
        if len(table_data) >= 2:
            tables.append(table_data)
    
    return '\n'.join(text_parts), tables


# main function that takes an uploaded file and returns structured syllabus data
def parse_syllabus(file):
    import os
    
    path = None
    
    try:
        # save the uploaded file to a temp location for processing
        with tempfile.NamedTemporaryFile(delete=False, suffix='.tmp') as tmp:
            tmp.write(file.file.read())
            path = tmp.name

        filename = file.filename.lower()
        text = ""
        tables = []

        # use different parser based on file type
        if filename.endswith(".pdf"):
            text, tables = extract_text_from_pdf(path)
        elif filename.endswith(".docx"):
            text, tables = extract_tables_from_docx(path)
        else:
            raise ValueError("Unsupported syllabus format. Please upload PDF or DOCX.")

        # try to find the subject code like "IC-812 Theory of Computation"
        subject = extract_subject_code(text)

        # first try to extract units from tables (most accurate)
        units = extract_units_from_tables(tables, text)

        # if no tables found, fall back to searching in plain text
        if not units:
            units = extract_units_from_text(text)

        # if still nothing found, just grab the first 50 meaningful lines
        if not units:
            lines = [line.strip() for line in text.splitlines() if len(line.strip()) > 10]
            units = [{
                "name": "General Topics",
                "title": "General Topics",
                "topics": lines[:50],
                "format": "long"
            }]

        return {
            "subject": subject,
            "units": units
        }
    
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(
            "Failed to read syllabus file. Please upload a valid PDF or DOCX."
        ) from e
    finally:
        # always clean up the temp file when done
        if path and os.path.exists(path):
            try:
                os.remove(path)
            except OSError:
                pass


# tries to find a subject code like "IC-812 Theory of Computation" in the text
def extract_subject_code(text: str) -> str:
    # patterns for different subject code formats
    subject_patterns = [
        r'([A-Z]{2,4}[-\s]?\d{3,4}[A-Z]?)\s*[:\-–]?\s*([A-Za-z][A-Za-z\s&,\-]+)',
        r'(?:Subject\s*Code|Course\s*Code|Code)\s*[:\-–]?\s*([A-Z]{2,4}[-\s]?\d{3,4}[A-Z]?)',
        r'(?:Subject\s*Name|Course\s*Name|Course\s*Title)\s*[:\-–]?\s*([A-Za-z][A-Za-z\s&,\-]+)',
    ]
    
    subject_code = ""
    subject_name = ""
    
    # try the most specific pattern first (code + name together)
    pattern1 = r'([A-Z]{2,4}[-\s]?\d{3,4}[A-Z]?)\s*[:\-–]?\s*([A-Za-z][A-Za-z\s&,\-]{3,50})'
    match = re.search(pattern1, text)
    if match:
        subject_code = match.group(1).strip()
        subject_name = match.group(2).strip()
        subject_name = re.sub(r'\s+', ' ', subject_name).strip()
        return f"{subject_code} {subject_name}"[:100]
    
    # try other patterns one by one
    for pattern in subject_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            result = match.group(1).strip()
            if match.lastindex >= 2:
                result = f"{result} {match.group(2).strip()}"
            return result[:100]
    
    # last resort: just use the first meaningful line as the subject
    for line in text.split('\n')[:10]:
        line = line.strip()
        if len(line) > 5 and not line.lower().startswith(('unit', 'module', 'chapter')):
            return line[:100]
    
    return "Unknown Subject"


# figures out which column is which in a table header row
def find_column_indices(header_row: List[str]) -> Dict[str, int]:
    indices = {"unit_no": -1, "title": -1, "contents": -1}
    
    # all the different names people use for these columns
    unit_names = ['unit no', 'unit', 'unit no.', 'module', 'module no', 'sl', 'sno', 's.no', 'no', 'sr']
    title_names = ['title', 'topic', 'unit title', 'module title', 'name', 'chapter', 'heading']
    content_names = ['contents', 'content', 'topics', 'syllabus', 'description', 'details', 'sub-topics', 'subtopics']
    
    # check each cell in the header to match column types
    for idx, cell in enumerate(header_row):
        cell_lower = cell.lower().strip()
        
        if any(name in cell_lower for name in unit_names) and indices["unit_no"] == -1:
            indices["unit_no"] = idx
        elif any(name in cell_lower for name in title_names) and indices["title"] == -1:
            indices["title"] = idx
        elif any(name in cell_lower for name in content_names) and indices["contents"] == -1:
            indices["contents"] = idx
    
    return indices


# splits a comma-separated list of topics from a table cell into individual items
def parse_contents_cell(content: str) -> List[str]:
    if not content:
        return []
    
    # split by comma, semicolon, or newline
    topics = re.split(r'[,;]\s*|\n+', content)
    
    # clean up each topic
    cleaned_topics = []
    for topic in topics:
        topic = topic.strip()
        # remove numbering like "1.", "a)", "(i)" etc
        topic = re.sub(r'^[\d]+[.\)]\s*|^[a-z][.\)]\s*|^\([ivxlc]+\)\s*', '', topic, flags=re.IGNORECASE)
        topic = topic.strip()
        
        # only keep topics with at least 3 characters
        if len(topic) >= 3:
            cleaned_topics.append(topic)
    
    return cleaned_topics


# extracts units and topics from table structures found in the document
def extract_units_from_tables(tables: List[List[List[str]]], text: str) -> List[Dict]:
    units = []
    
    for table in tables:
        # skip tables that are too small
        if len(table) < 2:
            continue
        
        # try to figure out which row is the header
        header_idx = 0
        indices = find_column_indices(table[0])
        
        # if first row doesnt look like a header, try the second row
        if indices["contents"] == -1 and len(table) > 1:
            indices = find_column_indices(table[1])
            if indices["contents"] != -1:
                header_idx = 1
        
        # skip this table if we cant find a contents column
        if indices["contents"] == -1:
            continue
        
        # go through each data row and extract unit info
        for row_idx, row in enumerate(table[header_idx + 1:], start=1):
            if len(row) <= max(i for i in indices.values() if i >= 0):
                continue
            
            # get the unit number from the row
            unit_no = ""
            if indices["unit_no"] >= 0 and indices["unit_no"] < len(row):
                unit_no = row[indices["unit_no"]].strip()
                # clean up the unit number
                unit_no = re.sub(r'^unit\s*', '', unit_no, flags=re.IGNORECASE).strip()
            
            if not unit_no:
                unit_no = str(row_idx)
            
            # get the unit title
            title = ""
            if indices["title"] >= 0 and indices["title"] < len(row):
                title = row[indices["title"]].strip()
            
            # get the topics from the contents column
            contents = ""
            if indices["contents"] >= 0 and indices["contents"] < len(row):
                contents = row[indices["contents"]]
            
            topics = parse_contents_cell(contents)
            
            # only add this unit if it actually has topics
            if topics:
                unit_name = f"Unit {unit_no}"
                if title:
                    unit_name = f"Unit {unit_no}: {title}"
                
                units.append({
                    "name": unit_name,
                    "title": title or f"Unit {unit_no}",
                    "topics": topics,
                    "format": infer_format(topics)
                })
    
    return units


# fallback: tries to find units by scanning plain text line by line
def extract_units_from_text(text: str) -> List[Dict]:
    units = []
    current_unit = None
    current_title = ""
    current_topics = []

    lines = text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # look for lines that start a new unit (like "Unit I: Introduction")
        unit_match = re.match(
            r'(?:unit|module|chapter)\s*([ivxlc0-9]+)\s*[:\-–]?\s*(.*)',
            line, re.IGNORECASE
        )

        if unit_match:
            # save the previous unit before starting a new one
            if current_unit and current_topics:
                units.append({
                    "name": current_unit,
                    "title": current_title,
                    "topics": current_topics,
                    "format": infer_format(current_topics)
                })

            # start tracking a new unit
            unit_num = unit_match.group(1).upper()
            current_title = unit_match.group(2).strip() if unit_match.group(2) else ""
            current_unit = f"Unit {unit_num}"
            if current_title:
                current_unit = f"Unit {unit_num}: {current_title}"
            current_topics = []

        # if we are inside a unit, add the line as a topic
        elif current_unit:
            # handle comma-separated topics on a single line
            if ',' in line and len(line) > 10:
                topics = parse_contents_cell(line)
                current_topics.extend(topics)
            elif len(line) > 4 and not re.match(r'^[\d\-•*]+$', line):
                current_topics.append(line)

    # dont forget to save the last unit
    if current_unit and current_topics:
        units.append({
            "name": current_unit,
            "title": current_title,
            "topics": current_topics,
            "format": infer_format(current_topics)
        })

    return units


# guesses whether topics need short, medium, or long answers based on their length
def infer_format(topics: List[str]) -> str:
    if not topics:
        return "short"

    # calculate the average length of all topics
    avg_length = sum(len(t) for t in topics) / len(topics)

    # longer topics usually need longer answers
    if avg_length > 100:
        return "long"
    elif avg_length > 40:
        return "medium"
    else:
        return "short"
